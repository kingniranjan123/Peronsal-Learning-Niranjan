"""
Spark In Action - DOCX v2 Final Generator
Uses matplotlib+networkx to render ALL diagrams locally.
No API calls. 100% diagram coverage guaranteed.
"""

import re, sys, io, time
from pathlib import Path
import matplotlib
matplotlib.use('Agg')  # Non-interactive backend
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
from matplotlib.patches import FancyBboxPatch, FancyArrowPatch
import networkx as nx
import numpy as np
from PIL import Image

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ─── Configuration ─────────────────────────────────────────────────────────
BASE_DIR     = Path(r"D:\Desktop\13th August 2023\python-output\python-inputs\a-process-telegram-uploads\Spark-In-Action")
CONCEPTS_DIR = BASE_DIR / "concepts"
OUTPUT_FILE  = BASE_DIR / "fules" / "Spark_In_Action_Study_Guide_v2.docx"
TEMP_IMG_DIR = BASE_DIR / "fules" / "_diagrams_local"
TEMP_IMG_DIR.mkdir(parents=True, exist_ok=True)

STRUCTURE = [
    ("part1-first-steps", "PART 1: FIRST STEPS", [
        ("chapter1-introduction-to-apache-spark", "Chapter 1: Introduction to Apache Spark"),
        ("chapter2-spark-fundamentals",           "Chapter 2: Spark Fundamentals"),
        ("chapter3-writing-spark-applications",   "Chapter 3: Writing Spark Applications"),
        ("chapter4-spark-api-in-depth",           "Chapter 4: The Spark API in Depth"),
    ]),
    ("part2-spark-family", "PART 2: MEET THE SPARK FAMILY", [
        ("chapter5-sparkling-queries-spark-sql",      "Chapter 5: Sparkling Queries with Spark SQL"),
        ("chapter6-ingesting-data-spark-streaming",   "Chapter 6: Ingesting Data with Spark Streaming"),
        ("chapter7-getting-smart-with-mllib",         "Chapter 7: Getting Smart with MLlib"),
        ("chapter8-ml-classification-and-clustering", "Chapter 8: ML - Classification and Clustering"),
        ("chapter9-connecting-dots-graphx",           "Chapter 9: Connecting the Dots with GraphX"),
    ]),
    ("part3-spark-ops", "PART 3: SPARK OPS", [
        ("chapter10-running-spark",             "Chapter 10: Running Spark"),
        ("chapter11-spark-standalone-cluster",  "Chapter 11: Running on a Spark Standalone Cluster"),
        ("chapter12-running-on-yarn-and-mesos", "Chapter 12: Running on YARN and Mesos"),
    ]),
    ("part4-bringing-it-together", "PART 4: BRINGING IT TOGETHER", [
        ("chapter13-realtime-dashboard",      "Chapter 13: Case Study - Real-Time Dashboard"),
        ("chapter14-deep-learning-spark-h2o", "Chapter 14: Deep Learning on Spark with H2O"),
    ]),
]

diagram_counter = [0]
ok_count   = [0]
fail_count = [0]

# ─── Colour palette ────────────────────────────────────────────────────────
PALETTE = {
    'navy':    '#1F497D', 'blue':    '#2E74B5', 'teal':    '#1F7A7A',
    'green':   '#27AE60', 'orange':  '#E67E22', 'red':     '#E74C3C',
    'purple':  '#8E44AD', 'yellow':  '#F39C12', 'gray':    '#95A5A6',
    'white':   '#FFFFFF', 'dark':    '#2D2D2D',
}

# ─── Mermaid → Graph Parser ────────────────────────────────────────────────
def parse_mermaid(code: str):
    """Parse simple graph TD/LR Mermaid into (nodes, edges, node_styles, title)."""
    lines = code.strip().split('\n')
    header = lines[0].strip() if lines else 'graph TD'
    direction = 'TB'
    if 'LR' in header: direction = 'LR'
    elif 'RL' in header: direction = 'RL'
    elif 'BT' in header: direction = 'BT'

    labels = {}
    colors = {}
    edges  = []
    font_colors = {}

    # Extract labels: A["Label text"] or A["Label\nmore text"]
    for pat in [r'(\w+)\["([^"]+)"\]', r"(\w+)\['([^']+)'\]",
                r'(\w+)\((.+?)\)', r'(\w+)\[([^\]]{1,80})\]']:
        for m in re.finditer(pat, code):
            nid, lbl = m.group(1), m.group(2)
            if nid not in labels:
                labels[nid] = lbl.replace('\\n', '\n')

    # Edges: A --> B, A -->|"label"| B, A --> B[text], A -.-> B
    edge_pat = re.compile(
        r'(\w+)\s*(?:-+>|-\.->|==>|--o|--x)\s*(?:\|"?([^"|]{0,40})"?\|\s*)?(\w+)'
    )
    for m in edge_pat.finditer(code):
        src, lbl, dst = m.group(1), (m.group(2) or '').strip(), m.group(3)
        edges.append((src, dst, lbl.replace('"', '')))
        if src not in labels: labels[src] = src
        if dst not in labels: labels[dst] = dst

    # Style: style NodeID fill:#color,color:#fff
    for m in re.finditer(r'style\s+(\w+)\s+fill:([#\w]+)', code):
        colors[m.group(1)] = m.group(2)
    for m in re.finditer(r'style\s+(\w+)\s+.*?color:([#\w]+)', code):
        font_colors[m.group(1)] = m.group(2)

    return labels, edges, colors, font_colors, direction


def render_mermaid_local(mermaid_code: str, index: int, title_hint: str = '') -> Path:
    """Render Mermaid diagram locally using matplotlib + networkx."""
    try:
        labels, edges, colors, font_colors, direction = parse_mermaid(mermaid_code)

        if not labels and not edges:
            return _fallback_text_diagram(mermaid_code, index, title_hint)

        # Build directed graph
        G = nx.DiGraph()
        for nid in labels:
            G.add_node(nid)
        edge_labels = {}
        for src, dst, lbl in edges:
            G.add_edge(src, dst)
            if lbl:
                edge_labels[(src, dst)] = lbl

        # Choose layout
        n = len(G.nodes)
        if n == 0:
            return _fallback_text_diagram(mermaid_code, index, title_hint)

        if direction in ('TB', 'TD') and n <= 15:
            try:
                pos = nx.nx_agraph.graphviz_layout(G, prog='dot')
            except Exception:
                pos = _hierarchical_layout(G)
        elif direction == 'LR' and n <= 15:
            try:
                pos = nx.nx_agraph.graphviz_layout(G, prog='dot', args='-Grankdir=LR')
            except Exception:
                pos = _left_right_layout(G)
        else:
            pos = nx.spring_layout(G, k=2.5, iterations=50, seed=42)

        # ── Draw ────────────────────────────────────────────────────────────
        fig_w = max(10, min(16, n * 1.4))
        fig_h = max(6,  min(12, n * 1.0))
        fig, ax = plt.subplots(1, 1, figsize=(fig_w, fig_h))
        ax.set_facecolor('#F8FAFC')
        fig.patch.set_facecolor('#F8FAFC')
        ax.axis('off')

        # Default node color if not styled
        default_colors = _auto_color_nodes(list(G.nodes), edges)

        node_colors = []
        node_font_colors = []
        for nid in G.nodes:
            bg = colors.get(nid, default_colors.get(nid, PALETTE['blue']))
            node_colors.append(bg)
            node_font_colors.append(font_colors.get(nid, '#FFFFFF'))

        # Draw edges first
        nx.draw_networkx_edges(
            G, pos, ax=ax,
            edge_color='#5A5A5A',
            arrows=True,
            arrowsize=20,
            arrowstyle='-|>',
            width=1.8,
            connectionstyle='arc3,rad=0.05',
            min_source_margin=25,
            min_target_margin=25,
        )

        # Draw edge labels
        if edge_labels:
            nx.draw_networkx_edge_labels(
                G, pos, edge_labels=edge_labels, ax=ax,
                font_size=7, font_color='#333333',
                bbox=dict(boxstyle='round,pad=0.2', facecolor='#FFFFF0', alpha=0.7),
            )

        # Draw nodes as boxes
        for nid, (x, y) in pos.items():
            bg   = colors.get(nid, default_colors.get(nid, PALETTE['blue']))
            fg   = font_colors.get(nid, '#FFFFFF')
            lbl  = labels.get(nid, nid)
            # Wrap long labels
            wrapped = _wrap_label(lbl, 18)
            lines   = wrapped.split('\n')
            box_w   = max(0.18, min(0.30, len(max(lines, key=len)) * 0.012))
            box_h   = 0.06 + 0.035 * len(lines)

            box = FancyBboxPatch(
                (x - box_w/2, y - box_h/2), box_w, box_h,
                boxstyle="round,pad=0.01",
                linewidth=1.5,
                edgecolor='#FFFFFF',
                facecolor=bg,
                transform=ax.transData,
                zorder=3,
            )
            ax.add_patch(box)
            ax.text(
                x, y, wrapped,
                ha='center', va='center',
                fontsize=7.5, color=fg,
                fontweight='bold',
                wrap=True,
                zorder=4,
                multialignment='center',
            )

        # Title
        display_title = title_hint.replace('_', ' ').replace('-', ' ').title()
        ax.set_title(display_title, fontsize=11, color=PALETTE['navy'],
                     fontweight='bold', pad=8)

        plt.tight_layout(pad=0.5)
        img_path = TEMP_IMG_DIR / f"diag_{index:03d}.png"
        plt.savefig(str(img_path), dpi=130, bbox_inches='tight',
                    facecolor=fig.get_facecolor())
        plt.close(fig)
        ok_count[0] += 1
        return img_path

    except Exception as e:
        fail_count[0] += 1
        print(f"    ! render error: {e}")
        return _fallback_text_diagram(mermaid_code, index, title_hint)


def _wrap_label(text: str, max_chars: int = 18) -> str:
    """Word-wrap a label to max_chars per line."""
    text = text.replace('\\n', '\n')
    lines = []
    for part in text.split('\n'):
        if len(part) <= max_chars:
            lines.append(part)
        else:
            words = part.split()
            cur = ''
            for w in words:
                if len(cur) + len(w) + 1 <= max_chars:
                    cur = (cur + ' ' + w).strip()
                else:
                    if cur: lines.append(cur)
                    cur = w
            if cur: lines.append(cur)
    return '\n'.join(lines[:4])  # max 4 lines per node


def _auto_color_nodes(nodes: list, edges: list) -> dict:
    """Auto-assign colours based on position in graph (first=blue, last=green, middle=teal)."""
    srcs = {e[0] for e in edges}
    dsts = {e[1] for e in edges}
    colors = {}
    for n in nodes:
        if n not in srcs:
            colors[n] = PALETTE['green']   # Sink / leaf
        elif n not in dsts:
            colors[n] = PALETTE['navy']    # Source / root
        else:
            colors[n] = PALETTE['teal']    # Middle
    return colors


def _hierarchical_layout(G: nx.DiGraph) -> dict:
    """Simple top-down hierarchical layout."""
    if len(G.nodes) == 0:
        return {}
    try:
        levels = {}
        roots = [n for n in G.nodes if G.in_degree(n) == 0]
        if not roots:
            roots = list(G.nodes)[:1]
        queue = list(roots)
        for r in roots: levels[r] = 0
        while queue:
            node = queue.pop(0)
            for nb in G.successors(node):
                if nb not in levels:
                    levels[nb] = levels[node] + 1
                    queue.append(nb)
        for n in G.nodes:
            if n not in levels: levels[n] = 0
        max_level = max(levels.values()) if levels else 0
        level_nodes = {}
        for n, l in levels.items():
            level_nodes.setdefault(l, []).append(n)
        pos = {}
        for l, ns in level_nodes.items():
            for i, n in enumerate(ns):
                x = (i - (len(ns)-1)/2) * 2.5
                y = (max_level - l) * 2.0
                pos[n] = (x, y)
        return pos
    except Exception:
        return nx.spring_layout(G, seed=42)


def _left_right_layout(G: nx.DiGraph) -> dict:
    """Simple left-to-right hierarchical layout."""
    topo = _hierarchical_layout(G)
    # Swap x and y for LR
    return {n: (y, -x) for n, (x, y) in topo.items()}


def _fallback_text_diagram(mermaid_code: str, index: int, title_hint: str = '') -> Path:
    """Create a styled text-based image when graph parsing fails."""
    fig, ax = plt.subplots(figsize=(10, 5))
    ax.set_facecolor('#EBF3FB')
    fig.patch.set_facecolor('#EBF3FB')
    ax.axis('off')

    title = title_hint.replace('_', ' ').replace('-', ' ').title()
    ax.text(0.5, 0.92, f"Architecture Diagram: {title}",
            ha='center', va='top', transform=ax.transAxes,
            fontsize=12, fontweight='bold', color=PALETTE['navy'])

    # Show first ~20 lines of Mermaid as styled text
    lines = mermaid_code.split('\n')[:20]
    display = '\n'.join(lines)
    ax.text(0.05, 0.82, display,
            ha='left', va='top', transform=ax.transAxes,
            fontsize=8, family='monospace', color='#333333',
            bbox=dict(boxstyle='round', facecolor='#FFFFFF', alpha=0.8, pad=0.5))

    img_path = TEMP_IMG_DIR / f"diag_{index:03d}.png"
    plt.savefig(str(img_path), dpi=100, bbox_inches='tight')
    plt.close(fig)
    fail_count[0] += 1
    return img_path


# ─── Colour helpers ─────────────────────────────────────────────────────────
def hex_to_rgb(h: str) -> RGBColor:
    h = h.lstrip('#')
    return RGBColor(int(h[0:2],16), int(h[2:4],16), int(h[4:6],16))

COLORS = {
    'navy':'#1F497D','blue':'#2E74B5','teal':'#1F7A7A',
    'mid_gray':'#595959','code_bg':'#F4F4F4','table_hdr':'#2E74B5',
    'table_row':'#EBF3FB','white':'#FFFFFF','accent':'#C55A11',
    'light_gray':'#AEAAAA',
}

def set_cell_bg(cell, hex_color: str):
    tcPr = cell._tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color.lstrip('#'))
    tcPr.append(shd)

def set_para_bg(p, hex_color: str):
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color.lstrip('#'))
    pPr.append(shd)

def add_left_border(p, color='2E74B5', sz='18'):
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bdr  = OxmlElement('w:left')
    bdr.set(qn('w:val'),'single'); bdr.set(qn('w:sz'),sz)
    bdr.set(qn('w:space'),'4');    bdr.set(qn('w:color'),color.lstrip('#'))
    pBdr.append(bdr); pPr.append(pBdr)

# ─── Document Setup ─────────────────────────────────────────────────────────
def setup_document() -> Document:
    doc = Document()
    for section in doc.sections:
        section.top_margin=Inches(1); section.bottom_margin=Inches(1)
        section.left_margin=Inches(1.25); section.right_margin=Inches(1)
    s = doc.styles
    n = s['Normal']; n.font.name='Calibri'; n.font.size=Pt(10.5)
    n.paragraph_format.space_after=Pt(6)
    n.paragraph_format.line_spacing_rule=WD_LINE_SPACING.MULTIPLE
    n.paragraph_format.line_spacing=1.15
    h1=s['Heading 1']; h1.font.name='Calibri Light'; h1.font.size=Pt(26); h1.font.bold=True; h1.font.color.rgb=hex_to_rgb('#FFFFFF')
    h2=s['Heading 2']; h2.font.name='Calibri Light'; h2.font.size=Pt(18); h2.font.bold=True; h2.font.color.rgb=hex_to_rgb(COLORS['navy'])
    h3=s['Heading 3']; h3.font.name='Calibri'; h3.font.size=Pt(13); h3.font.bold=True; h3.font.color.rgb=hex_to_rgb(COLORS['teal'])
    h4=s['Heading 4']; h4.font.name='Calibri'; h4.font.size=Pt(11); h4.font.bold=True; h4.font.color.rgb=hex_to_rgb(COLORS['blue'])
    return doc

def add_toc(doc):
    doc.add_heading('Table of Contents', level=2)
    note = doc.add_paragraph()
    r = note.add_run('Tip: In Word press Ctrl+A then F9 to generate the full Table of Contents.')
    r.font.italic=True; r.font.size=Pt(9); r.font.color.rgb=hex_to_rgb(COLORS['mid_gray'])
    set_para_bg(note,'#EBF3FB')
    para = doc.add_paragraph(); run = para.add_run()
    for tag, val in [('begin',''), ('separate',''), ('end','')]:
        fc = OxmlElement('w:fldChar'); fc.set(qn('w:fldCharType'), tag)
        if tag == 'begin':
            instr = OxmlElement('w:instrText'); instr.set(qn('xml:space'),'preserve')
            instr.text = 'TOC \\o "1-4" \\h \\z \\u'
            run._r.append(fc); run._r.append(instr)
        else:
            run._r.append(fc)
    doc.add_page_break()

def add_cover(doc):
    tbl = doc.add_table(1,1); tbl.alignment=WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.rows[0].cells[0]; set_cell_bg(cell, COLORS['navy'])
    p = cell.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before=Pt(36); p.paragraph_format.space_after=Pt(6)
    r = p.add_run('SPARK IN ACTION'); r.font.name='Calibri Light'; r.font.size=Pt(36)
    r.font.bold=True; r.font.color.rgb=hex_to_rgb('#FFFFFF')
    p2=cell.add_paragraph(); p2.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r2=p2.add_run('Complete Engineering Training Study Guide')
    r2.font.name='Calibri Light'; r2.font.size=Pt(17); r2.font.color.rgb=hex_to_rgb('#BDD7EE')
    p3=cell.add_paragraph(); p3.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r3=p3.add_run('v2  |  14 Chapters  |  91 Topics  |  4 Parts')
    r3.font.name='Calibri'; r3.font.size=Pt(11); r3.font.color.rgb=hex_to_rgb('#9DC3E6')
    p3.paragraph_format.space_after=Pt(36)
    doc.add_paragraph(); doc.add_page_break()

def add_part_banner(doc, title):
    tbl=doc.add_table(1,1); tbl.alignment=WD_TABLE_ALIGNMENT.CENTER
    cell=tbl.rows[0].cells[0]; set_cell_bg(cell,COLORS['navy'])
    p=cell.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before=Pt(12); p.paragraph_format.space_after=Pt(12)
    r=p.add_run(title); r.font.name='Calibri Light'; r.font.size=Pt(20)
    r.font.bold=True; r.font.color.rgb=hex_to_rgb('#FFFFFF')
    # Hidden heading for TOC
    h=doc.add_heading(title, level=1); h.clear()
    doc.add_paragraph()

def add_chapter_header(doc, title):
    p=doc.add_heading(title, level=2)
    pPr=p._p.get_or_add_pPr(); pBdr=OxmlElement('w:pBdr')
    bot=OxmlElement('w:bottom'); bot.set(qn('w:val'),'single')
    bot.set(qn('w:sz'),'8'); bot.set(qn('w:space'),'2')
    bot.set(qn('w:color'),COLORS['blue'].lstrip('#'))
    pBdr.append(bot); pPr.append(pBdr)

def format_inline(p, text):
    pat = re.compile(r'(\*\*(.+?)\*\*|\*(.+?)\*|`(.+?)`|__(.+?)__)')
    last=0
    for m in pat.finditer(text):
        if m.start()>last: p.add_run(text[last:m.start()])
        raw=m.group(0)
        if raw.startswith('**') or raw.startswith('__'):
            r=p.add_run(m.group(2) or m.group(5)); r.bold=True
        elif raw.startswith('`'):
            r=p.add_run(m.group(4)); r.font.name='Courier New'
            r.font.size=Pt(9); r.font.color.rgb=hex_to_rgb(COLORS['accent'])
        else:
            r=p.add_run(m.group(3)); r.italic=True
        last=m.end()
    if last<len(text): p.add_run(text[last:])

def embed_image(doc, img_path, caption):
    try:
        with Image.open(img_path) as img:
            w_px, h_px = img.size
        max_w = Inches(5.8)
        aspect = h_px/w_px
        pic_w = max_w
        pic_h = int(max_w*aspect)
        if pic_h > Inches(4.5):
            pic_h = Inches(4.5); pic_w = int(pic_h/aspect)
        p = doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(img_path), width=pic_w)
        cp=doc.add_paragraph(); cp.alignment=WD_ALIGN_PARAGRAPH.CENTER
        cr=cp.add_run(caption); cr.font.size=Pt(8.5); cr.font.italic=True
        cr.font.color.rgb=hex_to_rgb(COLORS['mid_gray'])
        cp.paragraph_format.space_after=Pt(8)
    except Exception as e:
        doc.add_paragraph(f'[Diagram — {caption}: {e}]')

def add_markdown_to_doc(doc, md_text, filename=''):
    lines = md_text.split('\n'); i=0
    while i < len(lines):
        line = lines[i]

        # Mermaid
        if line.strip().startswith('```mermaid'):
            mlines=[]
            i+=1
            while i<len(lines) and lines[i].strip()!='```':
                mlines.append(lines[i]); i+=1
            code='\n'.join(mlines).strip()
            diagram_counter[0]+=1
            idx=diagram_counter[0]
            hint=Path(filename).stem
            print(f'    >> Diagram #{idx} ({filename}) — local render')
            img=render_mermaid_local(code, idx, hint)
            if img and img.exists():
                embed_image(doc, img, f'Figure {idx}')
                print(f'       OK')
            else:
                doc.add_paragraph(f'[Diagram #{idx}]')
            i+=1; continue

        # Code block
        if line.strip().startswith('```'):
            lang=line.strip().replace('```','').strip(); clines=[]
            i+=1
            while i<len(lines) and lines[i].strip()!='```':
                clines.append(lines[i]); i+=1
            if lang:
                lp=doc.add_paragraph(); lr=lp.add_run(f' {lang.upper()} ')
                lr.font.name='Calibri'; lr.font.size=Pt(7.5); lr.font.bold=True
                lr.font.color.rgb=hex_to_rgb('#FFFFFF'); set_para_bg(lp,COLORS['teal'])
                lp.paragraph_format.space_after=Pt(0)
            p=doc.add_paragraph(); r=p.add_run('\n'.join(clines))
            r.font.name='Courier New'; r.font.size=Pt(8.5)
            set_para_bg(p,COLORS['code_bg']); add_left_border(p,'AAAAAA','6')
            p.paragraph_format.space_after=Pt(8); i+=1; continue

        # Table
        if line.strip().startswith('|') and i+1<len(lines) and lines[i+1].strip().startswith('|'):
            tlines=[]
            while i<len(lines) and lines[i].strip().startswith('|'):
                tlines.append(lines[i]); i+=1
            rows=[r for r in tlines if not re.match(r'^\|\s*[-:]+[\s|-]*\|?\s*$', r.strip())]
            if not rows: continue
            parsed=[]
            for rs in rows:
                cells=[c.strip() for c in rs.strip().strip('|').split('|')]
                parsed.append(cells)
            nc=max(len(r) for r in parsed)
            if nc==0: continue
            tbl=doc.add_table(len(parsed), nc); tbl.style='Table Grid'
            tbl.alignment=WD_TABLE_ALIGNMENT.CENTER
            for ri,rc in enumerate(parsed):
                while len(rc)<nc: rc.append('')
                for ci in range(nc):
                    cell=tbl.rows[ri].cells[ci]
                    txt=re.sub(r'\*+([^*]+)\*+',r'\1',rc[ci])
                    txt=re.sub(r'`([^`]+)`',r'\1',txt)
                    p=cell.paragraphs[0]; p.clear()
                    rn=p.add_run(txt); rn.font.size=Pt(9.5)
                    if ri==0:
                        rn.bold=True; rn.font.color.rgb=hex_to_rgb('#FFFFFF')
                        set_cell_bg(cell,COLORS['table_hdr'])
                    elif ri%2==0:
                        set_cell_bg(cell,COLORS['table_row'])
            doc.add_paragraph().paragraph_format.space_after=Pt(6); continue

        # Blockquote
        if line.strip().startswith('>'):
            txt=re.sub(r'^>\s*','',line.strip())
            if re.search(r'\[!(NOTE|TIP|IMPORTANT|WARNING|CAUTION)\]',txt):
                i+=1; continue
            p=doc.add_paragraph(); r=p.add_run(txt)
            r.font.size=Pt(10); r.font.italic=True
            r.font.color.rgb=hex_to_rgb(COLORS['mid_gray'])
            set_para_bg(p,'#F0F4FF'); add_left_border(p,COLORS['blue'],'18')
            p.paragraph_format.left_indent=Inches(0.15); i+=1; continue

        # Headings
        h4m=re.match(r'^####\s+(.*)',line)
        h3m=re.match(r'^###\s+(.*)',line)
        h2m=re.match(r'^##\s+(.*)',line)
        h1m=re.match(r'^#\s+(.*)',line)
        if h4m: doc.add_heading(h4m.group(1).strip(), level=4)
        elif h3m: doc.add_heading(h3m.group(1).strip(), level=4)
        elif h2m: doc.add_heading(h2m.group(1).strip(), level=3)
        elif h1m: doc.add_heading(h1m.group(1).strip(), level=3)

        # Lists
        elif re.match(r'^\s{0,3}[-*+]\s+',line):
            text=re.sub(r'^\s*[-*+]\s+','',line)
            indent=len(line)-len(line.lstrip())
            p=doc.add_paragraph(style='List Bullet 2' if indent>=2 else 'List Bullet')
            format_inline(p, text)
        elif re.match(r'^\s*\d+\.\s+',line):
            text=re.sub(r'^\s*\d+\.\s+','',line)
            p=doc.add_paragraph(style='List Number'); format_inline(p, text)

        # HR
        elif re.match(r'^-{3,}$|^={3,}$',line.strip()):
            p=doc.add_paragraph()
            pPr=p._p.get_or_add_pPr(); pBdr=OxmlElement('w:pBdr')
            bot=OxmlElement('w:bottom'); bot.set(qn('w:val'),'single')
            bot.set(qn('w:sz'),'4'); bot.set(qn('w:space'),'1')
            bot.set(qn('w:color'),COLORS['light_gray'].lstrip('#'))
            pBdr.append(bot); pPr.append(pBdr)

        # Normal text
        elif line.strip():
            p=doc.add_paragraph(); format_inline(p, line.strip())

        i+=1

def get_ordered_files(ch_path: Path):
    files=list(ch_path.glob('*.md'))
    overview=[f for f in files if 'overview' in f.name.lower()]
    topics=sorted([f for f in files if 'overview' not in f.name.lower()])
    return overview+topics

def main():
    print('='*65)
    print('  Spark In Action - DOCX Generator v2 (Local Render)')
    print('='*65)

    doc=setup_document()
    add_cover(doc)
    add_toc(doc)

    total=0
    for part_folder, part_title, chapters in STRUCTURE:
        part_path=CONCEPTS_DIR/part_folder
        if not part_path.exists():
            print(f'SKIP missing part: {part_folder}'); continue
        print(f'\n{"="*55}\n  {part_title}\n{"="*55}')
        doc.add_page_break()
        add_part_banner(doc, part_title)

        for ch_folder, ch_title in chapters:
            ch_path=part_path/ch_folder
            if not ch_path.exists():
                print(f'  SKIP: {ch_title}'); continue
            print(f'\n  [{ch_title}]')
            add_chapter_header(doc, ch_title)

            for md in get_ordered_files(ch_path):
                print(f'    {md.name}')
                try:
                    text=md.read_text(encoding='utf-8', errors='replace')
                    add_markdown_to_doc(doc, text, md.name)
                    doc.add_paragraph(); total+=1
                except Exception as e:
                    print(f'    ERROR: {e}')
            doc.add_page_break()

    OUTPUT_FILE.parent.mkdir(parents=True, exist_ok=True)
    print('\nSaving...')
    doc.save(str(OUTPUT_FILE))
    print('='*65)
    print(f'  COMPLETE!')
    print(f'  Files     : {total}')
    print(f'  Diagrams  : OK={ok_count[0]}  fallback={fail_count[0]}')
    print(f'  Output    : {OUTPUT_FILE}')
    print('='*65)

if __name__=='__main__':
    main()
