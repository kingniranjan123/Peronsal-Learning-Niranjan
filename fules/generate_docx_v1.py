"""
Spark in Action - Complete Study Guide DOCX Generator v1
Assembles all 92 MD files into a single, well-formatted DOCX document.
Mermaid diagrams are rendered via the mermaid.ink public API.
"""

import os
import re
import sys
import base64
import io
import time
import requests
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from PIL import Image

# ─────────────────────────────────────────────
# CONFIGURATION
# ─────────────────────────────────────────────
BASE_DIR = Path(r"D:\Desktop\13th August 2023\python-output\python-inputs\a-process-telegram-uploads\Spark-In-Action")
CONCEPTS_DIR = BASE_DIR / "concepts"
OUTPUT_FILE = BASE_DIR / "fules" / "Spark_In_Action_Study_Guide_v1.docx"
MERMAID_API = "https://mermaid.ink/img/{encoded}?type=png"
TEMP_IMG_DIR = BASE_DIR / "fules" / "_temp_diagrams"
TEMP_IMG_DIR.mkdir(parents=True, exist_ok=True)

# Ordered parts and chapters
STRUCTURE = [
    ("part1-first-steps", "Part 1: First Steps", [
        ("chapter1-introduction-to-apache-spark", "Chapter 1: Introduction to Apache Spark"),
        ("chapter2-spark-fundamentals",           "Chapter 2: Spark Fundamentals"),
        ("chapter3-writing-spark-applications",   "Chapter 3: Writing Spark Applications"),
        ("chapter4-spark-api-in-depth",           "Chapter 4: The Spark API in Depth"),
    ]),
    ("part2-spark-family", "Part 2: Meet the Spark Family", [
        ("chapter5-sparkling-queries-spark-sql",          "Chapter 5: Sparkling Queries with Spark SQL"),
        ("chapter6-ingesting-data-spark-streaming",       "Chapter 6: Ingesting Data with Spark Streaming"),
        ("chapter7-getting-smart-with-mllib",             "Chapter 7: Getting Smart with MLlib"),
        ("chapter8-ml-classification-and-clustering",     "Chapter 8: ML – Classification and Clustering"),
        ("chapter9-connecting-dots-graphx",               "Chapter 9: Connecting the Dots with GraphX"),
    ]),
    ("part3-spark-ops", "Part 3: Spark Ops", [
        ("chapter10-running-spark",          "Chapter 10: Running Spark"),
        ("chapter11-spark-standalone-cluster","Chapter 11: Running on a Spark Standalone Cluster"),
        ("chapter12-running-on-yarn-and-mesos","Chapter 12: Running on YARN and Mesos"),
    ]),
    ("part4-bringing-it-together", "Part 4: Bringing it Together", [
        ("chapter13-realtime-dashboard",           "Chapter 13: Case Study – Real-Time Dashboard"),
        ("chapter14-deep-learning-spark-h2o",      "Chapter 14: Deep Learning on Spark with H2O"),
    ]),
]

# ─────────────────────────────────────────────
# MERMAID RENDERING
# ─────────────────────────────────────────────
def render_mermaid(mermaid_code: str, index: int) -> Path | None:
    """Renders a Mermaid diagram via mermaid.ink and saves as PNG."""
    try:
        encoded = base64.urlsafe_b64encode(mermaid_code.encode("utf-8")).decode("utf-8")
        url = MERMAID_API.format(encoded=encoded)
        resp = requests.get(url, timeout=20)
        if resp.status_code == 200 and len(resp.content) > 500:
            img_path = TEMP_IMG_DIR / f"diagram_{index}.png"
            with open(img_path, "wb") as f:
                f.write(resp.content)
            # Verify it's a valid image
            try:
                img = Image.open(img_path)
                img.verify()
                return img_path
            except Exception:
                return None
        else:
            print(f"  ⚠ Mermaid API returned {resp.status_code} for diagram {index}")
            return None
    except Exception as e:
        print(f"  ⚠ Failed to render diagram {index}: {e}")
        return None


# ─────────────────────────────────────────────
# DOCUMENT STYLING
# ─────────────────────────────────────────────
def setup_document() -> Document:
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin    = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin   = Inches(1.2)
        section.right_margin  = Inches(1.2)

    styles = doc.styles

    # Normal
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    # Heading 1 — Part title
    h1 = styles["Heading 1"]
    h1.font.name = "Calibri Light"
    h1.font.size = Pt(26)
    h1.font.bold = True
    h1.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

    # Heading 2 — Chapter title
    h2 = styles["Heading 2"]
    h2.font.name = "Calibri Light"
    h2.font.size = Pt(20)
    h2.font.bold = True
    h2.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)

    # Heading 3 — Topic file title
    h3 = styles["Heading 3"]
    h3.font.name = "Calibri"
    h3.font.size = Pt(15)
    h3.font.bold = True
    h3.font.color.rgb = RGBColor(0x1F, 0x7A, 0x7A)

    # Heading 4 — Section within topic
    h4 = styles["Heading 4"]
    h4.font.name = "Calibri"
    h4.font.size = Pt(12)
    h4.font.bold = True
    h4.font.color.rgb = RGBColor(0x40, 0x40, 0x40)

    # Code style
    try:
        code_style = styles.add_style("CodeBlock", WD_STYLE_TYPE.PARAGRAPH)
    except Exception:
        code_style = styles["CodeBlock"]
    code_style.font.name = "Courier New"
    code_style.font.size = Pt(9)
    code_style.paragraph_format.space_before = Pt(4)
    code_style.paragraph_format.space_after  = Pt(4)

    return doc


# ─────────────────────────────────────────────
# COVER PAGE
# ─────────────────────────────────────────────
def add_cover(doc: Document):
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Spark in Action")
    run.font.name = "Calibri Light"
    run.font.size = Pt(40)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = subtitle.add_run("Complete Engineering Study Guide")
    run2.font.name = "Calibri Light"
    run2.font.size = Pt(22)
    run2.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)

    doc.add_paragraph()
    ver = doc.add_paragraph()
    ver.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run3 = ver.add_run("v1  |  14 Chapters  |  92 Topic Files  |  4 Parts")
    run3.font.name = "Calibri"
    run3.font.size = Pt(13)
    run3.font.color.rgb = RGBColor(0x70, 0x70, 0x70)

    doc.add_paragraph()
    tagline = doc.add_paragraph()
    tagline.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run4 = tagline.add_run("100-Day Production-Ready Data Engineer Bootcamp")
    run4.font.name = "Calibri"
    run4.font.size = Pt(13)
    run4.font.italic = True
    run4.font.color.rgb = RGBColor(0x40, 0x40, 0x40)

    doc.add_page_break()


# ─────────────────────────────────────────────
# TABLE OF CONTENTS PLACEHOLDER
# ─────────────────────────────────────────────
def add_toc(doc: Document):
    doc.add_heading("Table of Contents", level=1)
    p = doc.add_paragraph("(Right-click → Update Field to refresh the Table of Contents in Microsoft Word)")
    p.runs[0].font.italic = True
    p.runs[0].font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    # Insert Word TOC field
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    fldChar = OxmlElement("w:fldChar")
    fldChar.set(qn("w:fldCharType"), "begin")
    run._r.append(fldChar)

    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = 'TOC \\o "1-4" \\h \\z \\u'
    run._r.append(instrText)

    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "separate")
    run._r.append(fldChar2)

    fldChar3 = OxmlElement("w:fldChar")
    fldChar3.set(qn("w:fldCharType"), "end")
    run._r.append(fldChar3)

    doc.add_page_break()


# ─────────────────────────────────────────────
# MARKDOWN PARSER & DOCX WRITER
# ─────────────────────────────────────────────
diagram_counter = [0]

def add_markdown_to_doc(doc: Document, md_text: str):
    """Parse markdown and add content to the DOCX document."""
    lines = md_text.split("\n")
    i = 0
    while i < len(lines):
        line = lines[i]

        # ── Mermaid diagram block ──────────────────
        if line.strip().startswith("```mermaid"):
            mermaid_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith("```"):
                mermaid_lines.append(lines[i])
                i += 1
            mermaid_code = "\n".join(mermaid_lines)
            diagram_counter[0] += 1
            print(f"    🎨 Rendering Mermaid diagram #{diagram_counter[0]}...")
            img_path = render_mermaid(mermaid_code, diagram_counter[0])
            if img_path:
                try:
                    with Image.open(img_path) as img:
                        w, h = img.size
                    # Scale to fit page width (max 5.5 inches)
                    max_width = Inches(5.5)
                    aspect = h / w
                    width = max_width
                    height = int(max_width * aspect)
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = p.add_run()
                    run.add_picture(str(img_path), width=width)
                    cap = doc.add_paragraph(f"Figure {diagram_counter[0]}: Diagram")
                    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    cap.runs[0].font.size = Pt(9)
                    cap.runs[0].font.italic = True
                    cap.runs[0].font.color.rgb = RGBColor(0x70, 0x70, 0x70)
                    time.sleep(0.3)  # Throttle API
                except Exception as e:
                    p = doc.add_paragraph(f"[Diagram {diagram_counter[0]} — see source MD for Mermaid code]")
                    p.runs[0].font.italic = True
            else:
                # Fallback: show raw mermaid as code block
                p = doc.add_paragraph()
                run = p.add_run(f"[Diagram — Mermaid code]\n{mermaid_code}")
                run.font.name = "Courier New"
                run.font.size = Pt(8)
                run.font.color.rgb = RGBColor(0x50, 0x50, 0x50)
            i += 1
            continue

        # ── Generic code block ─────────────────────
        if line.strip().startswith("```"):
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith("```"):
                code_lines.append(lines[i])
                i += 1
            code_text = "\n".join(code_lines)
            # Add shaded code paragraph
            p = doc.add_paragraph()
            run = p.add_run(code_text)
            run.font.name = "Courier New"
            run.font.size = Pt(8.5)
            # Light gray background via XML shading
            pPr = p._p.get_or_add_pPr()
            shd = OxmlElement("w:shd")
            shd.set(qn("w:val"), "clear")
            shd.set(qn("w:color"), "auto")
            shd.set(qn("w:fill"), "F0F0F0")
            pPr.append(shd)
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after  = Pt(3)
            i += 1
            continue

        # ── Markdown Table ─────────────────────────
        if line.strip().startswith("|") and i + 1 < len(lines) and lines[i+1].strip().startswith("|"):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i])
                i += 1
            # Filter separator row
            rows = [r for r in table_lines if not re.match(r"^\|[\s\-\:]+\|", r.strip())]
            if not rows:
                continue
            cols = [c.strip() for c in rows[0].split("|") if c.strip()]
            num_cols = len(cols)
            if num_cols == 0:
                continue
            tbl = doc.add_table(rows=len(rows), cols=num_cols)
            tbl.style = "Table Grid"
            for r_idx, row_text in enumerate(rows):
                cells = [c.strip() for c in row_text.split("|") if c.strip()]
                # Pad or trim cells
                while len(cells) < num_cols:
                    cells.append("")
                for c_idx in range(num_cols):
                    cell = tbl.rows[r_idx].cells[c_idx]
                    # Strip markdown bold/italic from cell text
                    clean = re.sub(r"\*+([^*]+)\*+", r"\1", cells[c_idx])
                    cell.text = clean
                    if r_idx == 0:
                        for run in cell.paragraphs[0].runs:
                            run.font.bold = True
                            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                        # Header row shading
                        tc = cell._tc
                        tcPr = tc.get_or_add_tcPr()
                        shd = OxmlElement("w:shd")
                        shd.set(qn("w:val"), "clear")
                        shd.set(qn("w:color"), "auto")
                        shd.set(qn("w:fill"), "2E74B5")
                        tcPr.append(shd)
            doc.add_paragraph()
            continue

        # ── Headings ───────────────────────────────
        h4_match = re.match(r"^####\s+(.*)", line)
        h3_match = re.match(r"^###\s+(.*)", line)
        h2_match = re.match(r"^##\s+(.*)", line)
        h1_match = re.match(r"^#\s+(.*)", line)

        if h4_match:
            doc.add_heading(h4_match.group(1).strip(), level=4)
        elif h3_match:
            doc.add_heading(h3_match.group(1).strip(), level=3)
        elif h2_match:
            doc.add_heading(h2_match.group(1).strip(), level=2)
        elif h1_match:
            doc.add_heading(h1_match.group(1).strip(), level=3)

        # ── Bullet / numbered list ─────────────────
        elif re.match(r"^\s*[-*]\s+", line):
            clean = re.sub(r"^\s*[-*]\s+", "", line)
            clean = re.sub(r"\*\*(.+?)\*\*", r"\1", clean)
            clean = re.sub(r"`(.+?)`", r"\1", clean)
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(clean)

        elif re.match(r"^\s*\d+\.\s+", line):
            clean = re.sub(r"^\s*\d+\.\s+", "", line)
            clean = re.sub(r"\*\*(.+?)\*\*", r"\1", clean)
            clean = re.sub(r"`(.+?)`", r"\1", clean)
            p = doc.add_paragraph(style="List Number")
            p.add_run(clean)

        # ── Horizontal rule ────────────────────────
        elif re.match(r"^---+$", line.strip()):
            p = doc.add_paragraph()
            pPr = p._p.get_or_add_pPr()
            pBdr = OxmlElement("w:pBdr")
            bottom = OxmlElement("w:bottom")
            bottom.set(qn("w:val"), "single")
            bottom.set(qn("w:sz"), "6")
            bottom.set(qn("w:space"), "1")
            bottom.set(qn("w:color"), "2E74B5")
            pBdr.append(bottom)
            pPr.append(pBdr)

        # ── Normal paragraph ───────────────────────
        elif line.strip():
            # Inline formatting: **bold**, *italic*, `code`
            p = doc.add_paragraph()
            remaining = line.strip()
            pattern = re.compile(r"(\*\*(.+?)\*\*|\*(.+?)\*|`(.+?)`)")
            last_end = 0
            for m in pattern.finditer(remaining):
                if m.start() > last_end:
                    p.add_run(remaining[last_end:m.start()])
                if m.group(0).startswith("**"):
                    run = p.add_run(m.group(2))
                    run.bold = True
                elif m.group(0).startswith("*"):
                    run = p.add_run(m.group(3))
                    run.italic = True
                elif m.group(0).startswith("`"):
                    run = p.add_run(m.group(4))
                    run.font.name = "Courier New"
                    run.font.size = Pt(9.5)
                last_end = m.end()
            if last_end < len(remaining):
                p.add_run(remaining[last_end:])

        # ── Blank line ─────────────────────────────
        else:
            pass  # Skip blank lines (Word handles spacing via style)

        i += 1


# ─────────────────────────────────────────────
# FILE ORDER WITHIN A CHAPTER FOLDER
# ─────────────────────────────────────────────
def get_ordered_files(chapter_path: Path) -> list[Path]:
    """Returns files in order: overview first, then numbered topic files."""
    files = list(chapter_path.glob("*.md"))
    overview = [f for f in files if "overview" in f.name.lower()]
    topics   = sorted([f for f in files if "overview" not in f.name.lower()])
    return overview + topics


# ─────────────────────────────────────────────
# MAIN
# ─────────────────────────────────────────────
def main():
    print("=" * 60)
    print("  Spark in Action — DOCX Generator v1")
    print("=" * 60)

    doc = setup_document()
    add_cover(doc)
    add_toc(doc)

    total_files = 0

    for part_folder, part_title, chapters in STRUCTURE:
        part_path = CONCEPTS_DIR / part_folder
        if not part_path.exists():
            print(f"⚠ Part folder not found: {part_path}")
            continue

        print(f"\n📂 {part_title}")
        # Part-level heading with page break
        doc.add_page_break()
        p = doc.add_heading(part_title, level=1)

        for ch_folder, ch_title in chapters:
            ch_path = part_path / ch_folder
            if not ch_path.exists():
                print(f"  ⚠ Chapter folder not found: {ch_path}")
                continue

            print(f"  📖 {ch_title}")
            doc.add_heading(ch_title, level=2)

            ordered_files = get_ordered_files(ch_path)
            for md_file in ordered_files:
                print(f"    📄 {md_file.name}")
                try:
                    md_text = md_file.read_text(encoding="utf-8")
                    add_markdown_to_doc(doc, md_text)
                    doc.add_paragraph()  # Spacing between files
                    total_files += 1
                except Exception as e:
                    print(f"    ✗ Error reading {md_file.name}: {e}")

            doc.add_page_break()

    # Save
    OUTPUT_FILE.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUTPUT_FILE))
    print("\n" + "=" * 60)
    print(f"✅ DONE! {total_files} files assembled.")
    print(f"📄 Output: {OUTPUT_FILE}")
    print("=" * 60)


if __name__ == "__main__":
    main()
