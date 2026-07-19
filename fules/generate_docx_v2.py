"""
Spark in Action - Complete Study Guide DOCX Generator v2
- Uses kroki.io API for robust Mermaid diagram rendering (handles complex syntax)
- Significantly improved document formatting and typography
- Saves directly in project directory
- Better section dividers, callout boxes, and professional layout
"""

import os
import re
import sys
import base64
import zlib
import io
import time
import requests
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Emu, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from PIL import Image

# ─────────────────────────────────────────────────────────────
# CONFIGURATION
# ─────────────────────────────────────────────────────────────
BASE_DIR     = Path(r"D:\Desktop\13th August 2023\python-output\python-inputs\a-process-telegram-uploads\Spark-In-Action")
CONCEPTS_DIR = BASE_DIR / "concepts"
OUTPUT_FILE  = BASE_DIR / "fules" / "Spark_In_Action_Study_Guide_v2.docx"
TEMP_IMG_DIR = BASE_DIR / "fules" / "_temp_diagrams_v2"
TEMP_IMG_DIR.mkdir(parents=True, exist_ok=True)

# kroki.io supports full Mermaid syntax including subgraphs, sequence, etc.
KROKI_URL = "https://kroki.io/mermaid/png"

# Ordered structure
STRUCTURE = [
    ("part1-first-steps", "PART 1: FIRST STEPS", [
        ("chapter1-introduction-to-apache-spark", "Chapter 1: Introduction to Apache Spark"),
        ("chapter2-spark-fundamentals",           "Chapter 2: Spark Fundamentals"),
        ("chapter3-writing-spark-applications",   "Chapter 3: Writing Spark Applications"),
        ("chapter4-spark-api-in-depth",           "Chapter 4: The Spark API in Depth"),
    ]),
    ("part2-spark-family", "PART 2: MEET THE SPARK FAMILY", [
        ("chapter5-sparkling-queries-spark-sql",      "Chapter 5: Sparkling Queries with Spark SQL"),
        ("chapter6-ingesting-data-spark-streaming",   "Chapter 6: Ingesting Data with Spark Streaming"),
        ("chapter7-getting-smart-with-mllib",         "Chapter 7: Getting Smart with MLlib"),
        ("chapter8-ml-classification-and-clustering", "Chapter 8: ML - Classification and Clustering"),
        ("chapter9-connecting-dots-graphx",           "Chapter 9: Connecting the Dots with GraphX"),
    ]),
    ("part3-spark-ops", "PART 3: SPARK OPS", [
        ("chapter10-running-spark",             "Chapter 10: Running Spark"),
        ("chapter11-spark-standalone-cluster",  "Chapter 11: Running on a Spark Standalone Cluster"),
        ("chapter12-running-on-yarn-and-mesos", "Chapter 12: Running on YARN and Mesos"),
    ]),
    ("part4-bringing-it-together", "PART 4: BRINGING IT TOGETHER", [
        ("chapter13-realtime-dashboard",      "Chapter 13: Case Study - Real-Time Dashboard"),
        ("chapter14-deep-learning-spark-h2o", "Chapter 14: Deep Learning on Spark with H2O"),
    ]),
]

diagram_counter = [0]
success_count   = [0]
fail_count      = [0]

# ─────────────────────────────────────────────────────────────
# DIAGRAM RENDERING via kroki.io
# ─────────────────────────────────────────────────────────────
def render_mermaid_kroki(mermaid_code: str, index: int) -> Path | None:
    """Render Mermaid via kroki.io POST API - handles complex syntax."""
    try:
        headers = {"Content-Type": "text/plain", "Accept": "image/png"}
        resp = requests.post(KROKI_URL, data=mermaid_code.encode("utf-8"),
                             headers=headers, timeout=30)
        if resp.status_code == 200 and len(resp.content) > 200:
            img_path = TEMP_IMG_DIR / f"diagram_v2_{index}.png"
            img_path.write_bytes(resp.content)
            # Validate image
            try:
                with Image.open(img_path) as img:
                    img.verify()
                success_count[0] += 1
                return img_path
            except Exception:
                fail_count[0] += 1
                return None
        else:
            print(f"    ! Diagram {index}: kroki returned {resp.status_code}")
            fail_count[0] += 1
            return None
    except Exception as e:
        print(f"    ! Diagram {index} failed: {type(e).__name__}")
        fail_count[0] += 1
        return None


# ─────────────────────────────────────────────────────────────
# DOCUMENT SETUP & STYLING
# ─────────────────────────────────────────────────────────────
def hex_to_rgb(hex_str: str) -> RGBColor:
    h = hex_str.lstrip("#")
    return RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))

COLORS = {
    "navy":       "#1F497D",
    "blue":       "#2E74B5",
    "teal":       "#1F7A7A",
    "dark_gray":  "#2D2D2D",
    "mid_gray":   "#595959",
    "light_gray": "#AEAAAA",
    "white":      "#FFFFFF",
    "code_bg":    "#F4F4F4",
    "warn_bg":    "#FFF3CD",
    "info_bg":    "#D9EAF7",
    "success_bg": "#D9EAD3",
    "accent":     "#C55A11",
    "table_hdr":  "#2E74B5",
    "table_row":  "#EBF3FB",
}

def set_cell_bg(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color.lstrip("#"))
    tcPr.append(shd)

def set_para_bg(p, hex_color: str):
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color.lstrip("#"))
    pPr.append(shd)

def add_para_border(p, color: str = "2E74B5", side: str = "left", sz: str = "24"):
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bdr = OxmlElement(f"w:{side}")
    bdr.set(qn("w:val"), "single")
    bdr.set(qn("w:sz"), sz)
    bdr.set(qn("w:space"), "4")
    bdr.set(qn("w:color"), color.lstrip("#"))
    pBdr.append(bdr)
    pPr.append(pBdr)

def setup_document() -> Document:
    doc = Document()
    for section in doc.sections:
        section.top_margin    = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin   = Inches(1.25)
        section.right_margin  = Inches(1.0)

    s = doc.styles

    # Normal
    normal = s["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    normal.paragraph_format.line_spacing = 1.15

    # Heading 1 — Part divider
    h1 = s["Heading 1"]
    h1.font.name  = "Calibri Light"
    h1.font.size  = Pt(28)
    h1.font.bold  = True
    h1.font.color.rgb = hex_to_rgb(COLORS["white"])
    h1.paragraph_format.space_before = Pt(24)
    h1.paragraph_format.space_after  = Pt(12)

    # Heading 2 — Chapter title
    h2 = s["Heading 2"]
    h2.font.name  = "Calibri Light"
    h2.font.size  = Pt(18)
    h2.font.bold  = True
    h2.font.color.rgb = hex_to_rgb(COLORS["navy"])
    h2.paragraph_format.space_before = Pt(18)
    h2.paragraph_format.space_after  = Pt(8)

    # Heading 3 — Topic title
    h3 = s["Heading 3"]
    h3.font.name  = "Calibri"
    h3.font.size  = Pt(14)
    h3.font.bold  = True
    h3.font.color.rgb = hex_to_rgb(COLORS["teal"])
    h3.paragraph_format.space_before = Pt(14)
    h3.paragraph_format.space_after  = Pt(6)

    # Heading 4 — Section within topic
    h4 = s["Heading 4"]
    h4.font.name  = "Calibri"
    h4.font.size  = Pt(11.5)
    h4.font.bold  = True
    h4.font.color.rgb = hex_to_rgb(COLORS["blue"])
    h4.paragraph_format.space_before = Pt(10)
    h4.paragraph_format.space_after  = Pt(4)

    # Code block style
    try:
        cs = s.add_style("SparkCode", WD_STYLE_TYPE.PARAGRAPH)
    except Exception:
        cs = s["SparkCode"] if "SparkCode" in [x.name for x in s] else s["Normal"]
    cs.font.name = "Courier New"
    cs.font.size = Pt(8.5)
    cs.paragraph_format.space_before = Pt(2)
    cs.paragraph_format.space_after  = Pt(2)
    cs.paragraph_format.left_indent  = Inches(0.3)

    return doc


# ─────────────────────────────────────────────────────────────
# COVER PAGE
# ─────────────────────────────────────────────────────────────
def add_cover(doc: Document):
    # Full-width navy banner via table
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.rows[0].cells[0]
    set_cell_bg(cell, COLORS["navy"])
    cell.width = Inches(6.5)

    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(40)
    p.paragraph_format.space_after  = Pt(8)
    run = p.add_run("SPARK IN ACTION")
    run.font.name  = "Calibri Light"
    run.font.size  = Pt(38)
    run.font.bold  = True
    run.font.color.rgb = hex_to_rgb(COLORS["white"])

    p2 = cell.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run("Complete Engineering Training Study Guide")
    r2.font.name  = "Calibri Light"
    r2.font.size  = Pt(18)
    r2.font.color.rgb = hex_to_rgb("#BDD7EE")

    p3 = cell.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = p3.add_run("v2  |  14 Chapters  |  92 Topics  |  4 Parts")
    r3.font.name  = "Calibri"
    r3.font.size  = Pt(11)
    r3.font.color.rgb = hex_to_rgb("#9DC3E6")
    p3.paragraph_format.space_after = Pt(40)

    doc.add_paragraph()

    # Info cards
    cards = [
        ("PARTS",    "4", "Part 1: Core | Part 2: Spark Family | Part 3: Ops | Part 4: Projects"),
        ("CHAPTERS", "14", "From Spark Fundamentals to Deep Learning with H2O"),
        ("TOPICS",   "92", "One detailed file per concept with diagrams & code"),
        ("GOAL",     "100-Day", "Bootcamp to become a Production-Ready Data Engineer"),
    ]
    for label, value, desc in cards:
        tbl2 = doc.add_table(rows=1, cols=2)
        tbl2.alignment = WD_TABLE_ALIGNMENT.CENTER
        c1, c2 = tbl2.rows[0].cells[0], tbl2.rows[0].cells[1]
        set_cell_bg(c1, COLORS["blue"])
        c1.width = Inches(1.1)
        p = c1.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(f"{value}\n{label}")
        r.font.name = "Calibri"
        r.font.bold = True
        r.font.size = Pt(13)
        r.font.color.rgb = hex_to_rgb(COLORS["white"])
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after  = Pt(6)

        c2.width = Inches(5.4)
        p2 = c2.paragraphs[0]
        r2 = p2.add_run(desc)
        r2.font.name = "Calibri"
        r2.font.size = Pt(10)
        r2.font.color.rgb = hex_to_rgb(COLORS["dark_gray"])
        p2.paragraph_format.space_before = Pt(6)
        doc.add_paragraph()

    doc.add_page_break()


# ─────────────────────────────────────────────────────────────
# TABLE OF CONTENTS
# ─────────────────────────────────────────────────────────────
def add_toc(doc: Document):
    h = doc.add_heading("Table of Contents", level=2)
    note = doc.add_paragraph()
    r = note.add_run("Tip: Open in Microsoft Word, then press Ctrl+A, F9 to update all fields and generate the full linked TOC.")
    r.font.italic = True
    r.font.size   = Pt(9)
    r.font.color.rgb = hex_to_rgb(COLORS["mid_gray"])
    set_para_bg(note, "#EBF3FB")

    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    fldChar = OxmlElement("w:fldChar")
    fldChar.set(qn("w:fldCharType"), "begin")
    run._r.append(fldChar)
    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = 'TOC \\o "1-4" \\h \\z \\u'
    run._r.append(instrText)
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "separate")
    run._r.append(fldChar2)
    fldChar3 = OxmlElement("w:fldChar")
    fldChar3.set(qn("w:fldCharType"), "end")
    run._r.append(fldChar3)

    doc.add_page_break()


# ─────────────────────────────────────────────────────────────
# PART BANNER (styled heading)
# ─────────────────────────────────────────────────────────────
def add_part_banner(doc: Document, title: str):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.rows[0].cells[0]
    set_cell_bg(cell, COLORS["navy"])
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(14)
    run = p.add_run(title)
    run.font.name  = "Calibri Light"
    run.font.size  = Pt(22)
    run.font.bold  = True
    run.font.color.rgb = hex_to_rgb(COLORS["white"])
    # Add to heading outline (hidden heading for TOC)
    h = doc.add_heading(title, level=1)
    h.clear()
    doc.add_paragraph()


# ─────────────────────────────────────────────────────────────
# CHAPTER HEADER
# ─────────────────────────────────────────────────────────────
def add_chapter_header(doc: Document, title: str):
    p = doc.add_heading(title, level=2)
    # Blue bottom border
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:space"), "2")
    bottom.set(qn("w:color"), COLORS["blue"].lstrip("#"))
    pBdr.append(bottom)
    pPr.append(pBdr)


# ─────────────────────────────────────────────────────────────
# EMBED DIAGRAM IMAGE
# ─────────────────────────────────────────────────────────────
def embed_diagram(doc: Document, img_path: Path, caption: str):
    try:
        with Image.open(img_path) as img:
            w_px, h_px = img.size
        max_w = Inches(5.8)
        aspect = h_px / w_px
        pic_w = max_w
        pic_h = int(max_w * aspect)
        # Cap height
        if pic_h > Inches(4.5):
            pic_h = Inches(4.5)
            pic_w = int(pic_h / aspect)

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(str(img_path), width=pic_w)

        cap_p = doc.add_paragraph()
        cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap_run = cap_p.add_run(caption)
        cap_run.font.size   = Pt(8.5)
        cap_run.font.italic = True
        cap_run.font.color.rgb = hex_to_rgb(COLORS["mid_gray"])
        cap_p.paragraph_format.space_after = Pt(10)
    except Exception as e:
        doc.add_paragraph(f"[Diagram could not be embedded: {e}]")


# ─────────────────────────────────────────────────────────────
# CALLOUT / NOTE BOX
# ─────────────────────────────────────────────────────────────
def add_callout(doc: Document, text: str, kind: str = "info"):
    bg_map = {"info": "#D9EAF7", "warn": "#FFF3CD", "tip": "#D9EAD3", "code": "#F4F4F4"}
    icon_map = {"info": "INFO", "warn": "NOTE", "tip": "TIP", "code": "CODE"}
    bg = bg_map.get(kind, "#F4F4F4")
    icon = icon_map.get(kind, "NOTE")

    tbl = doc.add_table(rows=1, cols=2)
    c1, c2 = tbl.rows[0].cells[0], tbl.rows[0].cells[1]
    set_cell_bg(c1, COLORS["blue"] if kind == "info" else COLORS["accent"] if kind == "warn" else "#2E7D32")
    c1.width = Inches(0.45)
    p1 = c1.paragraphs[0]
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = p1.add_run(icon)
    r1.font.name = "Calibri"
    r1.font.size = Pt(7)
    r1.font.bold = True
    r1.font.color.rgb = hex_to_rgb(COLORS["white"])
    p1.paragraph_format.space_before = Pt(4)
    p1.paragraph_format.space_after  = Pt(4)

    set_cell_bg(c2, bg)
    p2 = c2.paragraphs[0]
    r2 = p2.add_run(text)
    r2.font.name = "Calibri"
    r2.font.size = Pt(9.5)
    p2.paragraph_format.space_before = Pt(4)
    p2.paragraph_format.space_after  = Pt(4)
    p2.paragraph_format.left_indent  = Inches(0.1)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


# ─────────────────────────────────────────────────────────────
# INLINE TEXT FORMATTER (bold, italic, code)
# ─────────────────────────────────────────────────────────────
def format_inline(p, text: str):
    pattern = re.compile(r"(\*\*(.+?)\*\*|\*(.+?)\*|`(.+?)`|__(.+?)__)")
    last = 0
    for m in pattern.finditer(text):
        if m.start() > last:
            p.add_run(text[last:m.start()])
        raw = m.group(0)
        if raw.startswith("**") or raw.startswith("__"):
            r = p.add_run(m.group(2) or m.group(5))
            r.bold = True
        elif raw.startswith("`"):
            r = p.add_run(m.group(4))
            r.font.name = "Courier New"
            r.font.size = Pt(9)
            r.font.color.rgb = hex_to_rgb(COLORS["accent"])
        else:
            r = p.add_run(m.group(3))
            r.italic = True
        last = m.end()
    if last < len(text):
        p.add_run(text[last:])


# ─────────────────────────────────────────────────────────────
# MARKDOWN -> DOCX PARSER
# ─────────────────────────────────────────────────────────────
def add_markdown_to_doc(doc: Document, md_text: str, filename: str = ""):
    lines = md_text.split("\n")
    i = 0
    while i < len(lines):
        line = lines[i]

        # ── Mermaid block ──────────────────────────────────────
        if line.strip().startswith("```mermaid"):
            mermaid_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip() == "```":
                mermaid_lines.append(lines[i])
                i += 1
            mermaid_code = "\n".join(mermaid_lines).strip()
            diagram_counter[0] += 1
            idx = diagram_counter[0]
            print(f"    >> Diagram #{idx} ({filename})")
            img_path = render_mermaid_kroki(mermaid_code, idx)
            if img_path:
                print(f"       OK - rendered")
                embed_diagram(doc, img_path, f"Figure {idx}: Architecture / Flow Diagram")
            else:
                # Render as styled text fallback
                print(f"       FALLBACK - code block")
                p_title = doc.add_paragraph()
                r = p_title.add_run("[Flow Diagram]")
                r.bold = True
                r.font.size = Pt(9)
                r.font.color.rgb = hex_to_rgb(COLORS["blue"])
                p = doc.add_paragraph()
                run = p.add_run(mermaid_code)
                run.font.name = "Courier New"
                run.font.size = Pt(8)
                run.font.color.rgb = hex_to_rgb("#333333")
                set_para_bg(p, "#F0F4FF")
                add_para_border(p, COLORS["blue"], "left", "12")
            i += 1
            time.sleep(0.5)
            continue

        # ── Generic code block ─────────────────────────────────
        if line.strip().startswith("```"):
            lang = line.strip().replace("```", "").strip()
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip() == "```":
                code_lines.append(lines[i])
                i += 1
            code_text = "\n".join(code_lines)
            if lang:
                lang_p = doc.add_paragraph()
                lr = lang_p.add_run(f" {lang.upper()} ")
                lr.font.name = "Calibri"
                lr.font.size = Pt(7.5)
                lr.font.bold = True
                lr.font.color.rgb = hex_to_rgb(COLORS["white"])
                set_para_bg(lang_p, COLORS["teal"])
                lang_p.paragraph_format.space_after = Pt(0)
            p = doc.add_paragraph()
            run = p.add_run(code_text)
            run.font.name = "Courier New"
            run.font.size = Pt(8.5)
            set_para_bg(p, COLORS["code_bg"])
            add_para_border(p, "#AAAAAA", "left", "6")
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after  = Pt(8)
            i += 1
            continue

        # ── Markdown Table ─────────────────────────────────────
        if line.strip().startswith("|") and i + 1 < len(lines) and lines[i+1].strip().startswith("|"):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i])
                i += 1
            rows = [r for r in table_lines
                    if not re.match(r"^\|\s*[-:]+[\s|-]*\|?\s*$", r.strip())]
            if not rows:
                continue
            parsed = []
            for row_str in rows:
                cells = [c.strip() for c in row_str.strip().strip("|").split("|")]
                parsed.append(cells)
            num_cols = max(len(r) for r in parsed) if parsed else 0
            if num_cols == 0:
                continue
            tbl = doc.add_table(rows=len(parsed), cols=num_cols)
            tbl.style = "Table Grid"
            tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
            for r_idx, row_cells in enumerate(parsed):
                while len(row_cells) < num_cols:
                    row_cells.append("")
                for c_idx in range(num_cols):
                    cell = tbl.rows[r_idx].cells[c_idx]
                    raw_text = re.sub(r"\*+([^*]+)\*+", r"\1", row_cells[c_idx])
                    raw_text = re.sub(r"`([^`]+)`", r"\1", raw_text)
                    p = cell.paragraphs[0]
                    p.clear()
                    run = p.add_run(raw_text)
                    run.font.size = Pt(9.5)
                    if r_idx == 0:
                        run.bold = True
                        run.font.color.rgb = hex_to_rgb(COLORS["white"])
                        set_cell_bg(cell, COLORS["table_hdr"])
                    elif r_idx % 2 == 0:
                        set_cell_bg(cell, COLORS["table_row"])
            doc.add_paragraph().paragraph_format.space_after = Pt(6)
            continue

        # ── Blockquote / callout ───────────────────────────────
        if line.strip().startswith(">"):
            text = re.sub(r"^>\s*", "", line.strip())
            # Detect GitHub-style alert
            if "[!NOTE]" in text or "[!TIP]" in text or "[!IMPORTANT]" in text or "[!WARNING]" in text:
                i += 1
                continue
            p = doc.add_paragraph()
            r = p.add_run(text)
            r.font.size   = Pt(10)
            r.font.italic = True
            r.font.color.rgb = hex_to_rgb(COLORS["mid_gray"])
            set_para_bg(p, "#F0F4FF")
            add_para_border(p, COLORS["blue"], "left", "18")
            p.paragraph_format.left_indent = Inches(0.15)
            i += 1
            continue

        # ── Horizontal rule ────────────────────────────────────
        if re.match(r"^-{3,}$", line.strip()) or re.match(r"^={3,}$", line.strip()):
            p = doc.add_paragraph()
            pPr = p._p.get_or_add_pPr()
            pBdr = OxmlElement("w:pBdr")
            bottom = OxmlElement("w:bottom")
            bottom.set(qn("w:val"), "single")
            bottom.set(qn("w:sz"), "4")
            bottom.set(qn("w:space"), "1")
            bottom.set(qn("w:color"), COLORS["light_gray"].lstrip("#"))
            pBdr.append(bottom)
            pPr.append(pBdr)
            p.paragraph_format.space_after = Pt(6)
            i += 1
            continue

        # ── Headings ───────────────────────────────────────────
        h4_m = re.match(r"^####\s+(.*)", line)
        h3_m = re.match(r"^###\s+(.*)", line)
        h2_m = re.match(r"^##\s+(.*)", line)
        h1_m = re.match(r"^#\s+(.*)", line)

        if h4_m:
            doc.add_heading(h4_m.group(1).strip(), level=4)
        elif h3_m:
            doc.add_heading(h3_m.group(1).strip(), level=4)
        elif h2_m:
            t = h2_m.group(1).strip()
            # Section headings get left-border accent treatment
            p = doc.add_heading(t, level=3)
        elif h1_m:
            p = doc.add_heading(h1_m.group(1).strip(), level=3)

        # ── Bullet lists ───────────────────────────────────────
        elif re.match(r"^\s{0,3}[-*+]\s+", line):
            indent_level = len(line) - len(line.lstrip())
            text = re.sub(r"^\s*[-*+]\s+", "", line)
            style = "List Bullet 2" if indent_level >= 2 else "List Bullet"
            p = doc.add_paragraph(style=style)
            format_inline(p, text)

        # ── Numbered lists ─────────────────────────────────────
        elif re.match(r"^\s*\d+\.\s+", line):
            text = re.sub(r"^\s*\d+\.\s+", "", line)
            p = doc.add_paragraph(style="List Number")
            format_inline(p, text)

        # ── Normal paragraph ───────────────────────────────────
        elif line.strip():
            p = doc.add_paragraph()
            format_inline(p, line.strip())

        i += 1


# ─────────────────────────────────────────────────────────────
# FILE ORDERING
# ─────────────────────────────────────────────────────────────
def get_ordered_files(chapter_path: Path) -> list[Path]:
    files   = list(chapter_path.glob("*.md"))
    overview = [f for f in files if "overview" in f.name.lower()]
    topics   = sorted([f for f in files if "overview" not in f.name.lower()])
    return overview + topics


# ─────────────────────────────────────────────────────────────
# MAIN
# ─────────────────────────────────────────────────────────────
def main():
    print("=" * 65)
    print("   Spark in Action - Study Guide DOCX Generator v2")
    print("=" * 65)

    doc = setup_document()
    add_cover(doc)
    add_toc(doc)

    total_files = 0

    for part_folder, part_title, chapters in STRUCTURE:
        part_path = CONCEPTS_DIR / part_folder
        if not part_path.exists():
            print(f"WARNING: Part folder missing: {part_path}")
            continue

        print(f"\n{'='*55}")
        print(f"  {part_title}")
        print(f"{'='*55}")

        doc.add_page_break()
        add_part_banner(doc, part_title)

        for ch_folder, ch_title in chapters:
            ch_path = part_path / ch_folder
            if not ch_path.exists():
                print(f"  SKIP (missing): {ch_title}")
                continue

            print(f"\n  [{ch_title}]")
            add_chapter_header(doc, ch_title)

            for md_file in get_ordered_files(ch_path):
                print(f"    Processing: {md_file.name}")
                try:
                    md_text = md_file.read_text(encoding="utf-8", errors="replace")
                    add_markdown_to_doc(doc, md_text, md_file.name)
                    doc.add_paragraph()
                    total_files += 1
                except Exception as e:
                    print(f"    ERROR: {md_file.name} -> {e}")

            doc.add_page_break()

    # Save to project directory
    OUTPUT_FILE.parent.mkdir(parents=True, exist_ok=True)
    print(f"\nSaving document...")
    doc.save(str(OUTPUT_FILE))

    print("\n" + "=" * 65)
    print(f"  COMPLETE!")
    print(f"  Files assembled : {total_files}")
    print(f"  Diagrams OK     : {success_count[0]}")
    print(f"  Diagrams fallback: {fail_count[0]}")
    print(f"  Output          : {OUTPUT_FILE}")
    print("=" * 65)


if __name__ == "__main__":
    main()
